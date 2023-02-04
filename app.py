import PyPDF2
import docx

def clean_text(text):
    # Reemplazar caracteres no compatibles con XML
    text = text.replace("\x00", " ")
    text = text.replace("\x0c", " ")
    return text

def pdf_to_word(pdf_file, word_file):
    # Leer el archivo PDF
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    
    # Crear un documento Word
    doc = docx.Document()
    
    # Iterar por cada página del PDF
    for page in pdf_reader.pages:
        # Extraer el texto de la página
        page_content = clean_text(page.extract_text())
        
        # Codificar el texto como Unicode
        page_content = page_content.encode("utf-8").decode("utf-8")
        
        # Agregar el texto a un nuevo párrafo en el documento Word
        doc.add_paragraph(page_content)
        
    # Guardar el documento Word
    doc.save(word_file)

def word_to_pdf(word_file, pdf_file):
    # Leer el archivo Word
    doc = docx.Document(word_file)
    
    # Crear un archivo PDF
    pdf_writer = PyPDF2.PdfFileWriter()
    
    # Iterar por cada párrafo en el documento Word
    for para in doc.paragraphs:
        # Crear una nueva página PDF
        pdf_page = PyPDF2.pdf.PageObject.createBlankPage(pdf_writer)
        
        # Escribir el texto del párrafo en la página PDF
        pdf_page.writeText(para.text)
        
        # Agregar la página PDF al archivo PDF
        pdf_writer.addPage(pdf_page)
    
    # Guardar el archivo PDF
    with open(pdf_file, 'wb') as fh:
        pdf_writer.write(fh)
